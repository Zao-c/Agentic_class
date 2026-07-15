import argparse
import csv
import hashlib
import json
import re
from pathlib import Path
from typing import Any, Dict, Iterable, List, Tuple

from docx import Document


PROJECT_ROOT = Path(__file__).resolve().parent.parent
QUESTION_ROOT = PROJECT_ROOT / "data" / "active" / "question_bank"
OUTPUT_ROOT = PROJECT_ROOT / "data" / "datasets" / "candidate"


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for block in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def clean(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip()


def dedup_key(question: str) -> str:
    normalized = re.sub(r"\W+", "", question).lower()
    return hashlib.sha256(normalized.encode("utf-8")).hexdigest()[:20]


def base_item(
    item_id: str,
    subtype: str,
    question: str,
    answer: str,
    source_path: Path,
    locator: str,
    excerpt: str,
) -> Dict[str, Any]:
    return {
        "item_id": item_id,
        "task_type": "knowledge_qa",
        "subtype": subtype,
        "input_turns": [{"role": "student", "content": question}],
        "reference_answer": answer,
        "grading_criteria": [],
        "required_slots": [],
        "expected_status": "completed",
        "expected_tools": ["course_retrieval"],
        "risk_level": "low",
        "equipment_scope": {
            "brand": "ABB",
            "model": None,
            "controller": None,
            "software_version": None,
            "scope_status": "course_context_only",
        },
        "source": {
            "relative_path": source_path.relative_to(PROJECT_ROOT).as_posix(),
            "sha256": sha256(source_path),
            "locator": locator,
            "excerpt": excerpt[:2000],
            "source_type": "course_question_bank",
        },
        "provenance_type": "source_extracted",
        "review": {
            "status": "needs_teacher_review",
            "reviewer_id": None,
            "reviewed_at": None,
            "decision": None,
            "note": None,
        },
        "privacy": {
            "status": "automated_no_personal_data_detected",
            "deidentification_method": "not_applicable",
            "retention_until": None,
            "deletion_status": "active",
        },
        "dataset": {
            "tier": "candidate",
            "version": "1.0.0",
            "split": "unassigned",
            "dedup_group": dedup_key(question),
            "leakage_group": source_path.name,
        },
    }


def extract_judgements(path: Path) -> List[Dict[str, Any]]:
    items = []
    paragraphs = [item.text.strip() for item in Document(path).paragraphs if item.text.strip()]
    for index, paragraph in enumerate(paragraphs, start=1):
        answer_match = re.search(r"（(正确|错误)）\s*$", paragraph)
        if not answer_match:
            continue
        number_match = re.match(r"\s*(\d+)\s*[、.]?\s*", paragraph)
        number = int(number_match.group(1)) if number_match else index
        question = paragraph[: answer_match.start()].strip()
        question = re.sub(r"^\s*\d+\s*[、.]?\s*", "", question)
        items.append(
            base_item(
                "qa_judgement_p%03d" % index,
                "judgement",
                question,
                answer_match.group(1),
                path,
                "paragraph:%d;question:%d" % (index, number),
                paragraph,
            )
        )
    return items


def extract_single_choice(path: Path) -> List[Dict[str, Any]]:
    items = []
    paragraphs = [item.text.strip() for item in Document(path).paragraphs if item.text.strip()]
    pairs: List[Tuple[str, str, int]] = []
    for index in range(0, len(paragraphs) - 1, 2):
        pairs.append((paragraphs[index], paragraphs[index + 1], index + 1))
    for fallback_number, (stem, options_text, paragraph_index) in enumerate(pairs, start=1):
        answer_match = re.search(r"[（(]\s*([A-D])\s*[）)]", stem)
        if not answer_match:
            continue
        number_match = re.match(r"\s*(\d+)\s*[、.]?\s*", stem)
        number = int(number_match.group(1)) if number_match else fallback_number
        answer_letter = answer_match.group(1)
        question_stem = re.sub(r"^\s*\d+\s*[、.]?\s*", "", stem)
        question_stem = question_stem[: answer_match.start()] + "（ ）" + question_stem[answer_match.end() :]
        option_lines = [clean(line) for line in options_text.splitlines() if clean(line)]
        options = {}
        for line in option_lines:
            match = re.match(r"([A-D])[、.．]\s*(.+)", line)
            if match:
                options[match.group(1)] = match.group(2)
        question = clean(question_stem) + "\n" + "\n".join(option_lines)
        answer = answer_letter
        if answer_letter in options:
            answer += "、" + options[answer_letter]
        items.append(
            base_item(
                "qa_single_choice_%03d" % number,
                "single_choice",
                question,
                answer,
                path,
                "paragraph:%d;question:%d" % (paragraph_index, number),
                stem + "\n" + options_text,
            )
        )
    return items


def task_blocks(path: Path) -> Dict[int, List[str]]:
    result: Dict[int, List[str]] = {}
    current: int | None = None
    for paragraph in [item.text.strip() for item in Document(path).paragraphs if item.text.strip()]:
        match = re.match(r"实训任务\s*(\d+)\s*(.*)", paragraph)
        if match:
            current = int(match.group(1))
            result[current] = [paragraph]
        elif current is not None:
            result[current].append(paragraph)
    return result


def extract_training_tasks(question_path: Path, answer_path: Path) -> List[Dict[str, Any]]:
    questions = task_blocks(question_path)
    answers = task_blocks(answer_path)
    items = []
    for number in sorted(set(questions) & set(answers)):
        answer_lines = answers[number]
        try:
            marker = next(index for index, line in enumerate(answer_lines) if line.startswith("参考答案"))
        except StopIteration:
            continue
        answer = clean(" ".join(answer_lines[marker + 1 :]))
        if not answer:
            continue
        question_lines = questions[number]
        question = clean(" ".join(question_lines[:4]))
        items.append(
            base_item(
                "qa_training_task_%03d" % number,
                "training_task",
                question,
                answer,
                answer_path,
                "training_task:%d" % number,
                clean(" ".join(answer_lines[:8])),
            )
        )
        items[-1]["source"]["question_relative_path"] = question_path.relative_to(
            PROJECT_ROOT
        ).as_posix()
        items[-1]["source"]["question_sha256"] = sha256(question_path)
    return items


def write_jsonl(path: Path, items: Iterable[Dict[str, Any]]) -> None:
    path.write_text(
        "".join(json.dumps(item, ensure_ascii=False) + "\n" for item in items),
        encoding="utf-8",
    )


def write_review_csv(path: Path, items: List[Dict[str, Any]]) -> None:
    with path.open("w", encoding="utf-8-sig", newline="") as stream:
        writer = csv.DictWriter(
            stream,
            fieldnames=[
                "item_id",
                "subtype",
                "question",
                "reference_answer",
                "source_path",
                "source_locator",
                "review_decision",
                "reviewer_id",
                "reviewer_role",
                "reviewed_at",
                "human_review_attestation",
                "source_verified",
                "privacy_checked",
                "safety_checked",
                "split",
                "reviewed_question",
                "reviewed_reference_answer",
                "reviewed_source_locator",
                "review_note",
            ],
        )
        writer.writeheader()
        for item in items:
            writer.writerow(
                {
                    "item_id": item["item_id"],
                    "subtype": item["subtype"],
                    "question": item["input_turns"][0]["content"],
                    "reference_answer": item["reference_answer"],
                    "source_path": item["source"]["relative_path"],
                    "source_locator": item["source"]["locator"],
                    "review_decision": "",
                    "reviewer_id": "",
                    "review_note": "",
                }
            )


def main() -> None:
    parser = argparse.ArgumentParser(description="从真实课程题库生成待教师审核的 QA 候选集")
    parser.add_argument("--output-root", type=Path, default=OUTPUT_ROOT)
    args = parser.parse_args()
    judgement_path = QUESTION_ROOT / "机器人编程(判断).docx"
    choice_path = QUESTION_ROOT / "机器人编程(单选).docx"
    task_path = QUESTION_ROOT / "工业机器人操作与编程实训任务题库.docx"
    answer_path = QUESTION_ROOT / "工业机器人操作与编程实训任务题库参考答案.docx"
    groups = {
        "judgement": extract_judgements(judgement_path),
        "single_choice": extract_single_choice(choice_path),
        "training_task": extract_training_tasks(task_path, answer_path),
    }
    items = [item for group in groups.values() for item in group]
    ids = [item["item_id"] for item in items]
    if len(ids) != len(set(ids)):
        raise ValueError("候选集 item_id 重复")
    args.output_root.mkdir(parents=True, exist_ok=True)
    write_jsonl(args.output_root / "course_qa_v1.jsonl", items)
    write_review_csv(args.output_root / "course_qa_v1_review.csv", items)
    manifest = {
        "schema_version": "1.0.0",
        "dataset_tier": "candidate",
        "review_status": "needs_teacher_review",
        "counts": {key: len(value) for key, value in groups.items()},
        "total": len(items),
        "metric_eligibility": False,
        "notice": "这些样本来自真实课程题库，但未经逐条教师审核，不计入正式评测指标。",
    }
    (args.output_root / "course_qa_v1_manifest.json").write_text(
        json.dumps(manifest, ensure_ascii=False, indent=2) + "\n", encoding="utf-8"
    )
    print(json.dumps(manifest, ensure_ascii=False))


if __name__ == "__main__":
    main()
